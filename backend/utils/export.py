import io
import re
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

def parse_markdown_to_docx(markdown_text: str) -> io.BytesIO:
    doc = docx.Document()
    
    # Split text into lines
    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Clean inline markdown markers for Word
        clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
        clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
        clean_line = re.sub(r'_(.*?)_', r'\1', clean_line)
        clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(clean_line[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(clean_line[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(clean_line[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(clean_line[5:], level=4)
        # Bullet list
        elif stripped.startswith("* ") or stripped.startswith("- "):
            doc.add_paragraph(clean_line[2:], style='List Bullet')
        # Numbered list
        elif stripped.split(".")[0].isdigit() and len(stripped.split(".")) > 1:
            parts = clean_line.split(".", 1)
            doc.add_paragraph(parts[1].strip(), style='List Number')
        # Regular paragraph
        else:
            doc.add_paragraph(clean_line)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def parse_markdown_to_pdf(markdown_text: str) -> io.BytesIO:
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom styles to prevent duplication and add nice spacing
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=22,
        leading=26,
        spaceAfter=12,
        alignment=TA_CENTER
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontSize=15,
        leading=18,
        spaceBefore=14,
        spaceAfter=6
    )
    h3_style = ParagraphStyle(
        'DocH3',
        parent=styles['Heading3'],
        fontSize=11,
        leading=14,
        spaceBefore=10,
        spaceAfter=4
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['BodyText'],
        fontSize=9.5,
        leading=13,
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        leftIndent=20,
        firstLineIndent=-10,
        fontSize=9.5,
        leading=13,
        spaceAfter=4
    )

    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Parse text content (simple sanitization for reportlab tags)
        clean_text = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        
        # Clean markdown formatting (*bold*, _italic_, etc.)
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean_text)
        clean_text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', clean_text)
        clean_text = re.sub(r'_(.*?)_', r'<i>\1</i>', clean_text)
        clean_text = re.sub(r'`(.*?)`', r'<font face="Courier">\1</font>', clean_text)
        
        if stripped.startswith("# "):
            story.append(Paragraph(clean_text[2:], title_style))
            story.append(Spacer(1, 10))
        elif stripped.startswith("## "):
            story.append(Paragraph(clean_text[3:], h2_style))
            story.append(Spacer(1, 6))
        elif stripped.startswith("### "):
            story.append(Paragraph(clean_text[4:], h3_style))
            story.append(Spacer(1, 4))
        elif stripped.startswith("* ") or stripped.startswith("- "):
            story.append(Paragraph(f"&bull; {clean_text[2:]}", bullet_style))
        elif stripped.split(".")[0].isdigit() and len(stripped.split(".")) > 1:
            parts = clean_text.split(".", 1)
            num = parts[0].strip()
            text = parts[1].strip()
            story.append(Paragraph(f"{num}. {text}", bullet_style))
        else:
            story.append(Paragraph(clean_text, body_style))
            story.append(Spacer(1, 4))
            
    doc.build(story)
    file_stream.seek(0)
    return file_stream
