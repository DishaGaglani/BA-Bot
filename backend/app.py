import json
import io
import re
import requests
import urllib3

import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

from fastapi import FastAPI, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from database import engine, SessionLocal
import models

# Disable SSL Warnings for self-signed certificates or proxy contexts
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Create database tables
models.Base.metadata.create_all(bind=engine)

app = FastAPI(title="BA Bot API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

PREDICTION_URL = "https://forjinn.com/api/v1/prediction/249fc96e-5b62-4208-8787-0d77367e9eaf"


class MessageRequest(BaseModel):
    message: str


class PredictionRequest(BaseModel):
    question: str
    sessionId: str | None = None
    projectId: int | None = None


class MessageResponse(BaseModel):
    status: str
    reply: str


class ProjectPayload(BaseModel):
    id: int | None = None
    sessionId: str | None = None
    pdfGenerated: bool | None = False
    messages: list[dict[str, object]] | None = None
    project: dict[str, str]
    overview: dict[str, object]
    discovery: dict[str, object]
    functional_requirements: list[dict[str, object]]
    missing_fields: list[str]
    next_question: str


def parse_markdown_to_docx(markdown_text: str) -> io.BytesIO:
    doc = docx.Document()
    
    # Split text into lines
    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Clean inline markdown markers for Word
        clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
        clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
        clean_line = re.sub(r'_(.*?)_', r'\1', clean_line)
        clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(clean_line[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(clean_line[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(clean_line[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(clean_line[5:], level=4)
        # Bullet list
        elif stripped.startswith("* ") or stripped.startswith("- "):
            doc.add_paragraph(clean_line[2:], style='List Bullet')
        # Numbered list
        elif stripped.split(".")[0].isdigit() and len(stripped.split(".")) > 1:
            parts = clean_line.split(".", 1)
            doc.add_paragraph(parts[1].strip(), style='List Number')
        # Regular paragraph
        else:
            doc.add_paragraph(clean_line)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def parse_markdown_to_pdf(markdown_text: str) -> io.BytesIO:
    file_stream = io.BytesIO()
    doc = SimpleDocTemplate(file_stream, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom styles to prevent duplication and add nice spacing
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=22,
        leading=26,
        spaceAfter=12,
        alignment=TA_CENTER
    )
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontSize=15,
        leading=18,
        spaceBefore=14,
        spaceAfter=6
    )
    h3_style = ParagraphStyle(
        'DocH3',
        parent=styles['Heading3'],
        fontSize=11,
        leading=14,
        spaceBefore=10,
        spaceAfter=4
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['BodyText'],
        fontSize=9.5,
        leading=13,
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        leftIndent=20,
        firstLineIndent=-10,
        fontSize=9.5,
        leading=13,
        spaceAfter=4
    )

    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Parse text content (simple sanitization for reportlab tags)
        clean_text = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        
        # Clean markdown formatting (*bold*, _italic_, etc.)
        clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean_text)
        clean_text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', clean_text)
        clean_text = re.sub(r'_(.*?)_', r'<i>\1</i>', clean_text)
        clean_text = re.sub(r'`(.*?)`', r'<font face="Courier">\1</font>', clean_text)
        
        if stripped.startswith("# "):
            story.append(Paragraph(clean_text[2:], title_style))
            story.append(Spacer(1, 10))
        elif stripped.startswith("## "):
            story.append(Paragraph(clean_text[3:], h2_style))
            story.append(Spacer(1, 6))
        elif stripped.startswith("### "):
            story.append(Paragraph(clean_text[4:], h3_style))
            story.append(Spacer(1, 4))
        elif stripped.startswith("* ") or stripped.startswith("- "):
            story.append(Paragraph(f"&bull; {clean_text[2:]}", bullet_style))
        elif stripped.split(".")[0].isdigit() and len(stripped.split(".")) > 1:
            parts = clean_text.split(".", 1)
            num = parts[0].strip()
            text = parts[1].strip()
            story.append(Paragraph(f"{num}. {text}", bullet_style))
        else:
            story.append(Paragraph(clean_text, body_style))
            story.append(Spacer(1, 4))
            
    doc.build(story)
    file_stream.seek(0)
    return file_stream


@app.get("/api/health")
def health_check() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/api/message", response_model=MessageResponse)
def receive_message(payload: MessageRequest) -> MessageResponse:
    return MessageResponse(
        status="received",
        reply=f"Received: {payload.message}",
    )


@app.get("/api/projects")
def list_projects(db: Session = Depends(get_db)):
    db_projects = db.query(models.Project).all()
    result = []
    updated = False
    for p in db_projects:
        try:
            proj_data = json.loads(p.data)
            if not proj_data.get("sessionId"):
                import uuid
                proj_data["sessionId"] = f"session-{uuid.uuid4()}"
                p.data = json.dumps(proj_data)
                updated = True
            proj_data["id"] = p.id
            result.append(proj_data)
        except Exception:
            pass
    if updated:
        db.commit()
    return result


@app.get("/api/project/{project_id}")
def get_project(project_id: int, db: Session = Depends(get_db)):
    db_project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Project not found")
    try:
        proj_data = json.loads(db_project.data)
        if not proj_data.get("sessionId"):
            import uuid
            proj_data["sessionId"] = f"session-{uuid.uuid4()}"
            db_project.data = json.dumps(proj_data)
            db.commit()
        proj_data["id"] = db_project.id
        return proj_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing project data: {str(e)}")


@app.post("/api/project")
def create_project(payload: ProjectPayload, db: Session = Depends(get_db)):
    if not payload.sessionId:
        import uuid
        payload.sessionId = f"session-{uuid.uuid4()}"
    project_data_json = json.dumps(payload.model_dump())
    db_project = models.Project(data=project_data_json)
    db.add(db_project)
    db.commit()
    db.refresh(db_project)
    try:
        proj_data = json.loads(db_project.data)
        proj_data["id"] = db_project.id
        return proj_data
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Error parsing created project: {str(e)}")


@app.put("/api/project/{project_id}")
def update_project(project_id: int, payload: ProjectPayload, db: Session = Depends(get_db)):
    db_project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Preserve existing sessionId if update payload lacks one
    try:
        existing_data = json.loads(db_project.data)
        existing_session_id = existing_data.get("sessionId")
        if existing_session_id and not payload.sessionId:
            payload.sessionId = existing_session_id
    except Exception:
        pass

    if not payload.sessionId:
        import uuid
        payload.sessionId = f"session-{uuid.uuid4()}"

    payload.id = project_id
    db_project.data = json.dumps(payload.model_dump())
    db.commit()
    try:
        proj_data = json.loads(db_project.data)
        proj_data["id"] = db_project.id
        return proj_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing updated project: {str(e)}")


@app.delete("/api/project/{project_id}")
def delete_project(project_id: int, db: Session = Depends(get_db)):
    db_project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Project not found")
    db.delete(db_project)
    db.commit()
    return {"status": "deleted"}


@app.get("/api/project/{project_id}/export")
def export_project(project_id: int, format: str, db: Session = Depends(get_db)):
    db_project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not db_project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    try:
        project_data = json.loads(db_project.data)
    except Exception:
        raise HTTPException(status_code=500, detail="Malformed project record")
        
    session_id = project_data.get("sessionId")
    project_info = project_data.get("project", {})
    project_name = project_info.get("name") if project_info else None
    if not project_name:
        project_name = f"Project_{project_id}"
    
    prompt = (
        f"The requirements interview discovery workshop is complete for project '{project_name}'. "
        "Please generate and compile the final, detailed, and polished Requirements Discovery Document (FDR) "
        "containing all project information, overview, stakeholders, business problem, business goals, and functional requirements. "
        "Format the output using clear Markdown headings, bullet points, and numbered lists."
    )
    
    chat_id = session_id or (f"project-{project_id}" if project_id else None)
    payload = {
        "question": prompt,
        "streaming": False
    }
    if chat_id:
        payload["chatId"] = chat_id
        payload["overrideConfig"] = {"sessionId": chat_id}
        
    try:
        response = requests.post(PREDICTION_URL, json=payload, timeout=180, verify=False)
        response.raise_for_status()
        res_data = response.json()
        
        document_text = res_data.get("text")
        if not document_text:
            output_obj = res_data.get("output")
            if isinstance(output_obj, dict):
                document_text = output_obj.get("content", "")
            elif isinstance(output_obj, str):
                document_text = output_obj
            else:
                document_text = ""
                
        if not document_text:
            raise HTTPException(status_code=500, detail="Prediction service returned empty compiled text")
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to generate compiled document from Forjinn flow: {str(e)}")
        
    if format.lower() == "docx":
        file_stream = parse_markdown_to_docx(document_text)
        filename = f"{project_name.replace(' ', '_')}_Requirements.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format.lower() == "pdf":
        file_stream = parse_markdown_to_pdf(document_text)
        filename = f"{project_name.replace(' ', '_')}_Requirements.pdf"
        media_type = "application/pdf"
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Supported: docx, pdf")
        
    return StreamingResponse(
        file_stream,
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )


@app.post("/api/predict")
def predict(payload: PredictionRequest, db: Session = Depends(get_db)) -> StreamingResponse:
    context_prefix = ""
    is_first_message = True
    if payload.projectId:
        db_project = db.query(models.Project).filter(models.Project.id == payload.projectId).first()
        if db_project:
            try:
                project_data = json.loads(db_project.data)
                
                messages = project_data.get("messages", [])
                user_msgs = [m for m in messages if m.get("role") == "user"]
                if len(user_msgs) > 0:
                    is_first_message = False
                    
                proj_info = project_data.get("project", {})
                if proj_info and proj_info.get("name"):
                    context_prefix = (
                        f"[Project Context - Name: {proj_info.get('name')}, "
                        f"Department: {proj_info.get('department')}, "
                        f"Sponsor: {proj_info.get('sponsor')}, "
                        f"Business Unit: {proj_info.get('business_unit')}, "
                        f"Expected Completion: {proj_info.get('expected_completion')}]\n"
                    )
            except Exception:
                pass

    question = payload.question
    if is_first_message and context_prefix:
        question = f"{context_prefix}User query: {payload.question}"

    payload_dict = {"question": question, "streaming": True}
    
    # Use sessionId or fallback to a project-specific identifier for native Forjinn memory
    chat_id = payload.sessionId or (f"project-{payload.projectId}" if payload.projectId else None)
    if chat_id:
        payload_dict["chatId"] = chat_id
        payload_dict["overrideConfig"] = {"sessionId": chat_id}

    def event_generator():
        try:
            response = requests.post(PREDICTION_URL, json=payload_dict, stream=True, timeout=90, verify=False)
            response.raise_for_status()
            for line in response.iter_lines():
                if line:
                    line_str = line.decode("utf-8", "ignore").strip()
                    if line_str.startswith("data:"):
                        yield f"{line_str}\n\n"
        except Exception as exc:
            err_data = json.dumps({"event": "error", "message": str(exc)})
            yield f"data: {err_data}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
